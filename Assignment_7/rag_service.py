import io
import os
import re
import tempfile
from pathlib import Path
import chromadb
import requests
import pypdf
import docx
from requests.exceptions import RequestException
from sentence_transformers import SentenceTransformer
from huggingface_hub.utils import disable_progress_bars

disable_progress_bars()

BASE_PATH = "E:/B TECH IT/Celebal Internship/Assignment_7"
ENV_PATH = Path(f"{BASE_PATH}/.env")

if ENV_PATH.exists():
    for line in ENV_PATH.read_text(encoding="utf-8").splitlines():
        if "=" in line and not line.startswith("#"):
            key, value = line.split("=", 1)
            os.environ.setdefault(key.strip(), value.strip().strip('"').strip("'"))

hf_token = os.getenv("HF_TOKEN")

# --- GLOBAL INITIALIZATIONS ---
EMBEDDING_MODEL = None

# Initialize ChromaDB globally so it stays persistent across pipeline calls
CHROMA_CLIENT = chromadb.Client()

def document_loader(file_path):
    file_path = Path(file_path)
    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")
    
    text = ""
    if file_path.suffix.lower() == ".pdf":
        document = pypdf.PdfReader(file_path)
        for page in document.pages:
            content = page.extract_text()
            if content:
                text += content + "\n"
    elif file_path.suffix.lower() == ".txt":
        text = file_path.read_text(encoding="utf-8")
    elif file_path.suffix.lower() == ".docx":
        document = docx.Document(str(file_path))
        for paragraph in document.paragraphs:
            text += paragraph.text + "\n"
    else:
        raise ValueError(f"Unsupported file type: {file_path.suffix}")
    return text

def chunk_text(text, chunk_size=1000, overlap=150):
    """
    Optimized chunk size. 2000 characters was too large for short documents,
    diluting context. 1000 provides much tighter, query-relevant context chunks.
    """
    chunks = []
    paragraphs = re.split(r'\n\s*\n', text)
    current_chunk = ""
    
    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
        if len(current_chunk) + len(para) > chunk_size:
            if current_chunk:
                chunks.append(current_chunk.strip())
                overlap_text = current_chunk[-overlap:]
                current_chunk = overlap_text + " " + para
            else:
                chunks.append(para)
                current_chunk = ""
        else:
            current_chunk = current_chunk + " " + para if current_chunk else para
            
    if current_chunk:
        chunks.append(current_chunk.strip())
    return chunks

def load_text_from_document(document):
    if isinstance(document, (str, Path)):
        file_path = Path(document)
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        text = ""
        total_pages = 1
        if file_path.suffix.lower() == ".pdf":
            pdf_reader = pypdf.PdfReader(file_path)
            total_pages = len(pdf_reader.pages)
            for page in pdf_reader.pages:
                content = page.extract_text()
                if content:
                    text += content + "\n"
        elif file_path.suffix.lower() == ".txt":
            text = file_path.read_text(encoding="utf-8")
            total_pages = 1
        elif file_path.suffix.lower() == ".docx":
            doc = docx.Document(str(file_path))
            text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
            total_pages = 1
        else:
            raise ValueError(f"Unsupported file type: {file_path.suffix}")
        return text, total_pages

    if hasattr(document, "read") and hasattr(document, "name"):
        suffix = Path(document.name).suffix.lower()
        content_bytes = document.read()
        if hasattr(document, "seek"):
            document.seek(0)
            
        text = ""
        total_pages = 1
        if suffix == ".pdf":
            pdf_reader = pypdf.PdfReader(io.BytesIO(content_bytes))
            total_pages = len(pdf_reader.pages)
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        elif suffix == ".txt":
            text = content_bytes.decode("utf-8", errors="ignore")
            total_pages = 1
        elif suffix == ".docx":
            doc = docx.Document(io.BytesIO(content_bytes))
            text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
            total_pages = 1
        else:
            raise ValueError("Unsupported document input type")
        return text, total_pages
    raise ValueError("Unsupported document input type")

def get_embedding_model():
    global EMBEDDING_MODEL
    if EMBEDDING_MODEL is None:
        print("Loading embedding model...")
        EMBEDDING_MODEL = SentenceTransformer("all-MiniLM-L6-v2")
    return EMBEDDING_MODEL

def create_lightweight_embedding(text, dimensions=384):
    tokens = re.findall(r"\w+", text.lower())
    if not tokens:
        return [0.0] * dimensions
    vector = [0.0] * dimensions
    for token in tokens:
        index = abs(hash(token)) % dimensions
        vector[index] += 1.0
    norm = sum(value * value for value in vector) ** 0.5
    if norm == 0:
        return [0.0] * dimensions
    return [value / norm for value in vector]

def embedding_generation(chunks):
    if hf_token:
        os.environ["HF_TOKEN"] = hf_token
    try:
        embedding_model = get_embedding_model()
        # Enforce list formatting for SentenceTransformer compatibility
        if isinstance(chunks, str):
            chunks = [chunks]
        return embedding_model.encode(chunks, show_progress_bar=False)
    except Exception as exc:
        print(f"Embedding model failed ({exc}). Using lightweight fallback embeddings.")
        if isinstance(chunks, str):
            chunks = [chunks]
        return [create_lightweight_embedding(chunk) for chunk in chunks]

def vector_store_creation(chunks, embeddings):
    collection = CHROMA_CLIENT.get_or_create_collection("document_embeddings")
    
    # Safely flush old entries to prevent mixed context from prior runs
    try:
        existing = collection.get()
        if existing and existing.get("ids"):
            collection.delete(ids=existing["ids"])
    except Exception:
        pass

    embedding_values = embeddings.tolist() if hasattr(embeddings, "tolist") else embeddings
    collection.add(
        ids=[f"chunk_{i}" for i in range(len(embedding_values))],
        embeddings=embedding_values,
        documents=chunks
    )
    return collection

def query_processing(query, vector_store):
    query_embedding = embedding_generation(query)
    # Ensure nested list structure format required by ChromaDB query engine
    if isinstance(query_embedding, list) and not isinstance(query_embedding[0], list):
        query_embeddings = [query_embedding]
    else:
        query_embeddings = query_embedding.tolist() if hasattr(query_embedding, "tolist") else query_embedding
        
    return vector_store.query(query_embeddings=query_embeddings, n_results=3)

def clean_context_text(text):
    cleaned = re.sub(r"[\r\n]+", " ", text)
    cleaned = re.sub(r"\s+", " ", cleaned).strip()
    cleaned = re.sub(r"\bPage\s*\d+\b", "", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"\bpp\.\s*\d+\b", "", cleaned, flags=re.IGNORECASE)
    return cleaned.strip()

def extract_context(results):
    documents = results.get("documents", [])
    if not documents:
        return ""
    flattened = []
    for item in documents:
        if isinstance(item, list):
            flattened.extend(str(sub) for sub in item if sub)
        elif item:
            flattened.append(str(item))
    return clean_context_text(" ".join(flattened))

def call_groq_api(prompt, temperature=0.0):
    api_key = os.getenv("GROQ_API_KEY")
    if not api_key:
        raise ValueError("GROQ_API_KEY not found in environment.")
        
    url = "https://api.groq.com/openai/v1/chat/completions"
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json"
    }
    
    models = ["llama-3.1-8b-instant", "llama3-8b-8192", "gemma2-9b-it"]
    last_err = None
    for model in models:
        try:
            payload = {
                "model": model,
                "messages": [
                    {"role": "user", "content": prompt}
                ],
                "temperature": temperature,
                "max_tokens": 1024
            }
            response = requests.post(url, headers=headers, json=payload, timeout=20)
            response.raise_for_status()
            res_json = response.json()
            return res_json["choices"][0]["message"]["content"].strip()
        except Exception as e:
            last_err = e
            print(f"Groq API error for model {model}: {e}")
            continue
    raise last_err

def generate_llm_response(prompt, temperature=0.0, timeout=30):
    # Tier 1: Try Groq API
    if os.getenv("GROQ_API_KEY"):
        try:
            return call_groq_api(prompt, temperature=temperature)
        except Exception as e:
            print(f"Groq API failed: {e}. Falling back to local Ollama.")
    else:
        print("GROQ_API_KEY not configured in environment. Trying local Ollama.")
        
    # Tier 2: Try local Ollama
    try:
        response = requests.post(
            "http://localhost:11434/api/generate",
            json={
                "model": "gemma:2b",
                "prompt": prompt,
                "stream": False,
                "options": {
                    "temperature": temperature
                }
            },
            timeout=timeout
        )
        response.raise_for_status()
        return response.json()["response"].strip()
    except Exception as e:
        raise RuntimeError(f"Ollama local service failed or not reachable: {e}")

def context_retrieval(query, context):
    """
    Routes context + query to LLM (Groq or Ollama) with a stricter system instruction format
    to prevent it from ignoring the document content.
    """
    prompt = f"""Instructions: You are a strict factual assistant. Answer the Question based ONLY on the provided Context. 
If the context does not contain the answer, reply exactly with: "I don't know based on the given information." 
Do not extrapolate or use outside knowledge.

Context: {context}

Question: {query}
Answer:"""

    try:
        return generate_llm_response(prompt, temperature=0.0, timeout=30)
    except Exception as e:
        print(f"LLM generation failed ({e}). Attempting internal fallback scoring algorithm.")
        return build_fallback_answer(query, context)

def build_fallback_answer(query, context):
    if not context:
        return "I don't have enough context to answer that question."
    cleaned_context = clean_context_text(context)
    sentences = re.split(r"(?<=[.!?])\s+", cleaned_context)
    query_terms = re.findall(r"\w+", query.lower())
    
    scored_sentences = []
    for sentence in sentences:
        if len(sentence) < 25:
            continue
        sentence_lower = sentence.lower()
        score = sum(1 for term in query_terms if term in sentence_lower)
        scored_sentences.append((score, sentence.strip()))
        
    scored_sentences.sort(key=lambda item: item[0], reverse=True)
    selected = [s for score, s in scored_sentences if score > 0][:3]
    return " ".join(selected) if selected else cleaned_context[:300]

# --- EXECUTION PIPELINE ---
def ingest_document(document):
    """
    Parses document, chunks it, generates embeddings, stores them in ChromaDB,
    and returns a metadata dictionary.
    """
    text, total_pages = load_text_from_document(document)
    if not text or not text.strip():
        raise ValueError("Unable to extract text from the uploaded document.")
        
    chunks = chunk_text(text, chunk_size=1000, overlap=150)
    if not chunks:
        raise ValueError("No text chunks generated from the document.")
        
    embeddings = embedding_generation(chunks)
    vector_store = vector_store_creation(chunks, embeddings)
    
    return {
        "total_pages": total_pages,
        "chunks": chunks,
        "collection_name": "document_embeddings",
        "text": text
    }

def answer_query(ingested_data, query):
    """
    Queries the vector store and gets context to retrieve the answer.
    """
    if not ingested_data:
        return "Please upload a document first."
        
    collection_name = ingested_data.get("collection_name", "document_embeddings")
    collection = CHROMA_CLIENT.get_or_create_collection(collection_name)
    results = query_processing(query, collection)
    context = extract_context(results)
    
    if not context.strip():
        return "I don't know based on the given information."
        
    return context_retrieval(query, context)

def build_fallback_summary(text):
    """
    Generates a heuristic summary if Ollama fails or is not available.
    """
    if not text:
        return "No text to summarize."
    cleaned_text = clean_context_text(text)
    sentences = re.split(r"(?<=[.!?])\s+", cleaned_text)
    sentences = [s.strip() for s in sentences if len(s.strip()) > 15]
    
    if not sentences:
        return text[:300] + "..."
        
    # Take intro and conclusion
    intro = sentences[:3]
    conclusion = sentences[-2:] if len(sentences) > 5 else []
    
    # Simple keyword extraction to identify key details
    all_words = re.findall(r"\b\w{5,}\b", cleaned_text.lower())
    from collections import Counter
    word_counts = Counter(all_words)
    top_words = [word for word, count in word_counts.most_common(5)]
    
    body_sentences = []
    if len(sentences) > 5:
        middle_sentences = sentences[3:-2]
        scored = []
        for s in middle_sentences:
            score = sum(1 for word in top_words if word in s.lower())
            scored.append((score, s))
        scored.sort(key=lambda x: x[0], reverse=True)
        body_sentences = [s for score, s in scored[:3] if score > 0]
        
    summary_parts = intro
    if body_sentences:
        summary_parts.append("\n\n**Key Details:**")
        summary_parts.extend(body_sentences)
    if conclusion:
        summary_parts.append("\n\n**Conclusion:**")
        summary_parts.extend(conclusion)
        
    return " ".join(summary_parts)

def summarize_document(ingested_data):
    """
    Summarizes the ingested document using LLM (Groq or Ollama), with a heuristic fallback.
    """
    if not ingested_data or not ingested_data.get("text"):
        return "No document text available to summarize."
        
    text = ingested_data["text"]
    
    # Truncate text if it is extremely long to prevent LLM issues
    max_summary_input_chars = 6000
    if len(text) > max_summary_input_chars:
        input_text = text[:4000] + "\n... [text truncated for summarization] ...\n" + text[-2000:]
    else:
        input_text = text

    prompt = f"""Instructions: Provide a concise, comprehensive summary of the following document. Highlight the main topics, key points, and overall conclusion.
    
Document:
{input_text}

Summary:"""

    try:
        return generate_llm_response(prompt, temperature=0.3, timeout=45)
    except Exception as e:
        print(f"LLM summarization failed ({e}). Using fallback extractive summary.")
        return build_fallback_summary(text)

# Maintain legacy pipeline alias for backward compatibility just in case
def pipeline(document, query=None):
    ingested = ingest_document(document)
    if query is None:
        return ingested
    return answer_query(ingested, query)

